"""DOCX creation tool."""

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from run_agent.config.constants import MIME_DOCX
from run_agent.tools.base import BaseTool
from run_agent.tools.documents._common import filename_with_ext, finalize


class CreateDocxTool(BaseTool):
    name = "create_docx"
    description = "Create a Microsoft Word document (.docx) with formatted content."

    def parameters_schema(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Document title"},
                "sections": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "heading": {"type": "string"},
                            "content": {"type": "string"},
                            "level": {"type": "integer", "default": 1},
                        },
                        "required": ["heading", "content"],
                    },
                    "description": "Document sections with headings and content",
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename (without extension)",
                    "default": "document",
                },
            },
            "required": ["title", "sections"],
        }

    async def execute(
        self,
        title: str,
        sections: list[dict],
        filename: str = "document",
        _context: dict | None = None,
        **_kwargs: Any,
    ) -> str:
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for section in sections:
            doc.add_heading(section["heading"], level=int(section.get("level", 1)))
            for paragraph in section["content"].split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        return await finalize(
            filename_with_ext(filename, "docx"),
            buffer.getvalue(),
            MIME_DOCX,
            _context,
        )
